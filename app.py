from flask import Flask, jsonify, send_from_directory
from flask_cors import CORS
import os

# =========================
# IMPORTACIONES BACKEND
# =========================
from backend.routes.auth import auth_bp
from backend.routes.actividades import actividades_bp
from backend.routes.usuarios import usuarios_bp
from backend.routes.historial import historial_bp
from backend.routes.reportes import reportes_bp
from backend.routes.ia import ia_bp
from backend.config import SECRET_KEY, DEBUG, PORT
from backend.database import init_db
from ia.ollama_client import verificar_conexion

from flask import request, send_file
from docx import Document
import io

# =========================
# CONFIGURACIÓN BASE
# =========================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PAGES_DIR = os.path.join(BASE_DIR, 'fronted', 'pages')


def create_app():
    """
    Crea y configura la aplicación Flask.
    """
    app = Flask(
        __name__,
        static_folder=os.path.join(BASE_DIR, 'fronted'),    # Carpeta raíz de estáticos (css, js, img)
        static_url_path='/fronted'     # URL pública: /fronted/css/... /fronted/js/...
    )
    app.secret_key = SECRET_KEY
    app.config["SECRET_KEY"] = SECRET_KEY

    # CORS solo para las rutas de API
    CORS(app, resources={r"/api/*": {"origins": "*"}})

    # =========================
    # REGISTRO DE BLUEPRINTS
    # =========================
    app.register_blueprint(auth_bp,        url_prefix='/api/auth')
    app.register_blueprint(actividades_bp, url_prefix='/api')
    app.register_blueprint(usuarios_bp,    url_prefix='/api')
    app.register_blueprint(historial_bp,   url_prefix='/api')
    app.register_blueprint(reportes_bp,    url_prefix='/api')
    app.register_blueprint(ia_bp,          url_prefix='/api')

    # =========================
    # ENDPOINT DE ESTADO
    # =========================
    @app.route('/api/status')
    def status():
        return jsonify({
            'db':      'Conectado',
            'ollama':  'Activo' if verificar_conexion() else 'Sin conexión',
            'version': '1.0.0',
        })

    # =========================
    # RUTAS FRONTEND
    # =========================

    # --- Raíz y Login ---
    @app.route('/')
    @app.route('/login')
    @app.route('/login/')
    def login_page():
        return send_from_directory(
            os.path.join(PAGES_DIR, 'login'),
            'index.html'
        )

    # --- Nueva Actividad (anidada dentro de login/) ---
    @app.route('/nueva-actividad')
    @app.route('/nueva-actividad/')
    def nueva_actividad_page():
        return send_from_directory(
            os.path.join(PAGES_DIR, 'login', 'nueva-actividad'),
            'index.html'
        )

    # --- Dashboard ---
    @app.route('/dashboard')
    @app.route('/dashboard/')
    def dashboard_page():
        return send_from_directory(
            os.path.join(PAGES_DIR, 'dashboard'),
            'index.html'
        )

    # --- Actividades ---
    @app.route('/actividades')
    @app.route('/actividades/')
    def actividades_page():
        return send_from_directory(
            os.path.join(PAGES_DIR, 'actividades'),
            'index.html'
        )

    # --- Búsqueda Avanzada ---
    @app.route('/busqueda-avanzada')
    @app.route('/busqueda-avanzada/')
    def busqueda_avanzada_page():
        return send_from_directory(
            os.path.join(PAGES_DIR, 'busqueda-avanzada'),
            'index.html'
        )

    # --- Configuración ---
    @app.route('/configuracion')
    @app.route('/configuracion/')
    def configuracion_page():
        return send_from_directory(
            os.path.join(PAGES_DIR, 'configuracion'),
            'index.html'
        )

    # --- Generador IA ---
    @app.route('/generador-ia')
    @app.route('/generador-ia/')
    def generador_ia_page():
        return send_from_directory(
            os.path.join(PAGES_DIR, 'generador-ia'),
            'index.html'
        )

    # --- Historial ---
    @app.route('/historial')
    @app.route('/historial/')
    def historial_page():
        return send_from_directory(
            os.path.join(PAGES_DIR, 'historial'),
            'index.html'
        )

    # --- Reportes ---
    @app.route('/reportes')
    @app.route('/reportes/')
    def reportes_page():
        return send_from_directory(
            os.path.join(PAGES_DIR, 'reportes'),
            'index.html'
        )

    # --- Usuarios ---
    @app.route('/usuarios')
    @app.route('/usuarios/')
    def usuarios_page():
        return send_from_directory(
            os.path.join(PAGES_DIR, 'usuarios'),
            'index.html'
        )

    @app.route('/export/word', methods=['POST'])
    @app.route('/export/word/', methods=['POST'])
    def export_word():

        data = request.json

        doc = Document("plantilla_informe.docx")

        def replace_text(doc, key, value):
            for p in doc.paragraphs:
                if key in p.text:
                    p.text = p.text.replace(key, value)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if key in p.text:
                                p.text = p.text.replace(key, value)

        replace_text(doc, "{{mes}}", data.get("mes", ""))
        replace_text(doc, "{{periodo}}", data.get("periodo", ""))
        replace_text(doc, "{{fecha_inicio}}", data.get("fecha_inicio", ""))
        replace_text(doc, "{{fecha_fin}}", data.get("fecha_fin", ""))
        replace_text(doc, "{{tipologia}}", data.get("tipologia", ""))
        replace_text(doc, "{{modalidad}}", data.get("modalidad", ""))
        replace_text(doc, "{{programa}}", data.get("programa", ""))
        replace_text(doc, "{{nombre}}", data.get("nombre", ""))
        replace_text(doc, "{{descripcion}}", data.get("descripcion", ""))
        replace_text(doc, "{{objetivo}}", data.get("objetivo", ""))
        replace_text(doc, "{{num_participantes}}", str(data.get("num_participantes", 0)))
        replace_text(doc, "{{horas_dedicadas}}", str(data.get("horas_dedicadas", 0)))
        replace_text(doc, "{{recursos}}", data.get("recursos", ""))
        replace_text(doc, "{{resultados}}", data.get("resultados", ""))
        replace_text(doc, "{{observaciones}}", data.get("observaciones", ""))

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name="Informe_Actividad_Extension.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # =========================
    # MANEJADOR DE ERRORES 404
    # =========================
    @app.errorhandler(404)
    def not_found(e):
        if request.path.startswith("/api/"):
            return jsonify({"error": "Endpoint no encontrado"}), 404

        return send_from_directory(
            os.path.join(PAGES_DIR, 'login'),
            'index.html'
        )

    return app


# =========================
# EJECUCIÓN PRINCIPAL
# =========================
if __name__ == '__main__':
    init_db()
    app = create_app()

    print(f"\n{'='*45}")
    print(f"  AcademiX — Servidor iniciado")
    print(f"{'='*45}")
    print(f"  URL:    http://localhost:{PORT}")
    print(f"  Debug:  {DEBUG}")
    print(f"  BD:     lista ✓")

    app.run(host='0.0.0.0', port=PORT, debug=DEBUG)